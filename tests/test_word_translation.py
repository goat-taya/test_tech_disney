from docx import Document

from translator.word import translate_word


def test_word_replaces_french_text_and_preserves_basic_structure_and_bold(tmp_path):
    source = tmp_path / "memo.docx"
    output = tmp_path / "memo_en.docx"

    document = Document()
    paragraph = document.add_paragraph()
    run = paragraph.add_run("Memo confidentiel")
    run.bold = True
    document.add_paragraph("Verifier les stocks", style="List Bullet")
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "Equipe"
    table.cell(0, 1).text = "Message interne"
    document.save(source)

    translate_word(source, output)

    translated = Document(output)

    assert translated.paragraphs[0].text == "[EN] Memo confidentiel"
    assert translated.paragraphs[0].runs[0].bold is True
    assert translated.paragraphs[1].text == "[EN] Verifier les stocks"
    assert translated.paragraphs[1].style.name == "List Bullet"
    assert translated.tables[0].cell(0, 0).text == "[EN] Equipe"
    assert translated.tables[0].cell(0, 1).text == "[EN] Message interne"
