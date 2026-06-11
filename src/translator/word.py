from pathlib import Path

from docx import Document

from translator.service import TranslationService


def _translate_paragraph(paragraph, service: TranslationService) -> None:
    if not paragraph.text.strip():
        return

    if paragraph.runs:
        first_run = paragraph.runs[0]
        first_run.text = service.translate_text(paragraph.text)
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(service.translate_text(paragraph.text))


def translate_word(
    input_path: str | Path,
    output_path: str | Path,
    service: TranslationService | None = None,
) -> Path:
    input_path = Path(input_path)
    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    service = service or TranslationService()

    document = Document(input_path)

    for paragraph in document.paragraphs:
        _translate_paragraph(paragraph, service)

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    _translate_paragraph(paragraph, service)

    document.save(output_path)
    return output_path
