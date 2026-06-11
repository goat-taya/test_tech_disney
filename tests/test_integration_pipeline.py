from docx import Document
from openpyxl import Workbook, load_workbook
from pypdf import PdfReader
from reportlab.pdfgen import canvas

from translator.pipeline import translate_directory


def test_translate_directory_processes_excel_word_and_pdf_end_to_end(tmp_path):
    input_dir = tmp_path / "input"
    output_dir = tmp_path / "output"
    input_dir.mkdir()

    workbook = Workbook()
    sheet = workbook.active
    sheet.append(["sku", "description", "category"])
    sheet.append(["A-001", "Description commerciale", "retail"])
    workbook.save(input_dir / "products.xlsx")

    document = Document()
    run = document.add_paragraph().add_run("Memo operationnel")
    run.bold = True
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "Equipe"
    table.cell(0, 1).text = "Action"
    document.save(input_dir / "memo.docx")

    pdf = canvas.Canvas(str(input_dir / "brochure.pdf"))
    pdf.drawString(72, 720, "Texte brochure")
    pdf.save()

    translated_paths = translate_directory(input_dir, output_dir)

    assert translated_paths == [
        output_dir / "brochure_en.pdf",
        output_dir / "memo_en.docx",
        output_dir / "products_en.xlsx",
    ]

    rows = list(load_workbook(output_dir / "products_en.xlsx").active.iter_rows(values_only=True))
    assert rows[0] == ("sku", "description", "description_en", "category")
    assert rows[1][2] == "[EN] Description commerciale"

    translated_doc = Document(output_dir / "memo_en.docx")
    assert translated_doc.paragraphs[0].text == "[EN] Memo operationnel"
    assert translated_doc.paragraphs[0].runs[0].bold is True
    assert translated_doc.tables[0].cell(0, 1).text == "[EN] Action"

    pdf_text = "\n".join(
        page.extract_text() or "" for page in PdfReader(str(output_dir / "brochure_en.pdf")).pages
    )
    assert "[EN] Texte brochure" in pdf_text
